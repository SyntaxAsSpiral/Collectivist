#!/usr/bin/env python3
"""
Documents Scanner Plugin
Scans document collections and extracts rich metadata including PDF properties, Word document metadata, and text analysis.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import yaml

# Document metadata extraction libraries
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from plugin_interface import CollectionScanner, CollectionItem, PluginRegistry


class DocumentsScanner(CollectionScanner):
    """Scanner for Obsidian vault collections."""

    def get_name(self) -> str:
        return "documents"

    def get_supported_types(self) -> List[str]:
        return ["file"]

    def get_categories(self) -> List[str]:
        return [
            "research_papers",
            "business_docs",
            "legal_documents",
            "educational_materials",
            "technical_docs",
            "personal_docs",
            "reports_presentations",
            "utilities_misc"
        ]

    def detect(self, path: Path) -> bool:
        """
        Detect if this path contains a document collection.
        Looks for common document file types.
        """
        if not path.is_dir():
            return False

        # Document file extensions
        doc_extensions = ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt', '.md', '.tex']

        # Count document files
        doc_files = []
        for ext in doc_extensions:
            doc_files.extend(list(path.glob(f'**/*{ext}')))

        # Require at least 5 document files to consider it a document collection
        return len(doc_files) >= 5

    def scan(self, root_path: Path, config: Dict[str, Any]) -> List[CollectionItem]:
        """
        Scan Obsidian vault for markdown files.

        Config options:
        - exclude_hidden: bool (default True) - exclude files starting with '.'
        - exclude_patterns: list - additional patterns to exclude
        - preserve_data: dict - existing descriptions/categories to preserve
        """
        exclude_hidden = config.get('exclude_hidden', True)
        exclude_patterns = config.get('exclude_patterns', [])
        preserve_data = config.get('preserve_data', {})

        items = []

        # Default document collection exclusions
        default_exclusions = [
            '.git/',
            '.DS_Store',
            'Thumbs.db',
            '__pycache__/',
            'node_modules/',
            '.obsidian/',  # Avoid mixing with Obsidian vaults
        ]

        all_exclusions = default_exclusions + exclude_patterns

        # Document file extensions
        doc_extensions = ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt', '.md', '.tex']

        # Find all document files
        for root, dirs, files in os.walk(root_path):
            root_path_obj = Path(root)

            # Skip excluded directories
            dirs[:] = [d for d in dirs if not any(
                pattern in str(root_path_obj / d) for pattern in all_exclusions
            )]

            for file in files:
                file_path = root_path_obj / file

                # Check if it's a document file
                if file_path.suffix.lower() not in doc_extensions:
                    continue

                # Skip hidden files if configured
                if exclude_hidden and file.startswith('.'):
                    continue

                # Skip excluded files
                if any(pattern in str(file_path) for pattern in all_exclusions):
                    continue

                # Get file stats
                stat = file_path.stat()

                # Extract document-specific metadata
                document_metadata = self._extract_document_metadata(file_path)

                # Preserve existing description/category if available
                existing = preserve_data.get(str(file_path), {})

                # Create item
                item = CollectionItem(
                    short_name=file_path.stem,
                    type="file",
                    size=stat.st_size,
                    created=datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    modified=datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    accessed=datetime.fromtimestamp(stat.st_atime).isoformat(),
                    path=str(file_path),
                    description=existing.get('description'),
                    category=existing.get('category'),
                    metadata={
                        'file_extension': file_path.suffix.lower(),
                        'document_metadata': document_metadata,
                        'word_count': document_metadata.get('word_count', 0),
                        'page_count': document_metadata.get('page_count', 0),
                        'author': document_metadata.get('author', ''),
                        'title': document_metadata.get('title', ''),
                        'has_text_content': document_metadata.get('has_text_content', False),
                    }
                )

                items.append(item)

        # Sort by modification time (most recent first)
        items.sort(key=lambda x: x.modified, reverse=True)

        return items

    def _extract_document_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document-specific metadata from various file types."""
        metadata = {}
        file_ext = file_path.suffix.lower()

        try:
            if file_ext in ['.txt', '.md', '.tex']:
                # Text-based documents
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                metadata.update(self._extract_text_metadata(content))
            elif file_ext == '.pdf':
                # PDF documents - basic file info for now
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_ext in ['.doc', '.docx']:
                # Word documents - basic file info for now
                metadata.update(self._extract_office_metadata(file_path))
            else:
                # Other document types - basic file info
                metadata['has_text_content'] = False
                metadata['word_count'] = 0
                metadata['page_count'] = 0
        except Exception:
            # If metadata extraction fails, provide basic info
            metadata['has_text_content'] = False
            metadata['word_count'] = 0
            metadata['page_count'] = 0

        return metadata

    def _extract_text_metadata(self, content: str) -> Dict[str, Any]:
        """Extract metadata from text-based documents."""
        metadata = {}

        # Basic content analysis
        metadata['has_text_content'] = True
        metadata['word_count'] = len(content.split())
        metadata['char_count'] = len(content)
        metadata['line_count'] = len(content.splitlines())

        # Try to extract title (first heading or first line)
        lines = content.splitlines()
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and not line.startswith('#'):  # Skip markdown headers for now
                metadata['title'] = line[:100]  # First non-empty line as title
                break

        # Check for markdown headers
        if content.startswith('#'):
            first_line = content.split('\n', 1)[0]
            metadata['title'] = first_line.lstrip('#').strip()[:100]

        return metadata

    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files using PyPDF2."""
        metadata = {}

        try:
            if not PYPDF2_AVAILABLE:
                # Fallback to basic file info
                metadata['has_text_content'] = True  # Assume PDFs have text
                metadata['word_count'] = 0
                metadata['page_count'] = 0
                metadata['author'] = ''
                metadata['title'] = file_path.stem
                return metadata

            # Open PDF with PyPDF2
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Basic PDF info
                metadata['page_count'] = len(pdf_reader.pages)
                metadata['has_text_content'] = True
                
                # Extract metadata from PDF info
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    metadata['title'] = pdf_info.get('/Title', file_path.stem) or file_path.stem
                    metadata['author'] = pdf_info.get('/Author', '') or ''
                    metadata['subject'] = pdf_info.get('/Subject', '') or ''
                    metadata['creator'] = pdf_info.get('/Creator', '') or ''
                    metadata['producer'] = pdf_info.get('/Producer', '') or ''
                    metadata['creation_date'] = pdf_info.get('/CreationDate', '') or ''
                    metadata['modification_date'] = pdf_info.get('/ModDate', '') or ''
                else:
                    metadata['title'] = file_path.stem
                    metadata['author'] = ''
                    metadata['subject'] = ''
                    metadata['creator'] = ''
                    metadata['producer'] = ''
                    metadata['creation_date'] = ''
                    metadata['modification_date'] = ''
                
                # Extract text content for word count (first few pages only for performance)
                text_content = ''
                max_pages_for_text = min(5, len(pdf_reader.pages))  # Limit to first 5 pages
                
                for page_num in range(max_pages_for_text):
                    try:
                        page = pdf_reader.pages[page_num]
                        text_content += page.extract_text() + '\n'
                    except Exception:
                        continue  # Skip pages that can't be extracted
                
                # Calculate word count from extracted text
                if text_content.strip():
                    metadata['word_count'] = len(text_content.split())
                    metadata['char_count'] = len(text_content)
                    
                    # Extract first meaningful line as title if no title in metadata
                    if metadata['title'] == file_path.stem:
                        lines = text_content.strip().split('\n')
                        for line in lines[:10]:  # Check first 10 lines
                            line = line.strip()
                            if line and len(line) > 5:  # Skip very short lines
                                metadata['title'] = line[:100]  # Limit title length
                                break
                else:
                    metadata['word_count'] = 0
                    metadata['char_count'] = 0

        except Exception as e:
            # If PDF processing fails, provide basic info
            metadata['has_text_content'] = False
            metadata['word_count'] = 0
            metadata['page_count'] = 0
            metadata['author'] = ''
            metadata['title'] = file_path.stem
            metadata['error'] = str(e)

        return metadata

    def _extract_office_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Office documents using python-docx."""
        metadata = {}
        file_ext = file_path.suffix.lower()

        try:
            if file_ext in ['.docx'] and PYTHON_DOCX_AVAILABLE:
                # Use python-docx for .docx files
                doc = Document(file_path)
                
                # Document properties
                core_props = doc.core_properties
                metadata['title'] = core_props.title or file_path.stem
                metadata['author'] = core_props.author or ''
                metadata['subject'] = core_props.subject or ''
                metadata['keywords'] = core_props.keywords or ''
                metadata['category'] = core_props.category or ''
                metadata['comments'] = core_props.comments or ''
                metadata['created'] = str(core_props.created) if core_props.created else ''
                metadata['modified'] = str(core_props.modified) if core_props.modified else ''
                metadata['last_modified_by'] = core_props.last_modified_by or ''
                
                # Extract text content
                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                full_text = '\n'.join(text_content)
                metadata['has_text_content'] = bool(full_text.strip())
                metadata['word_count'] = len(full_text.split()) if full_text else 0
                metadata['char_count'] = len(full_text)
                metadata['paragraph_count'] = len([p for p in doc.paragraphs if p.text.strip()])
                
                # Estimate page count (rough approximation: 250 words per page)
                metadata['page_count'] = max(1, metadata['word_count'] // 250) if metadata['word_count'] > 0 else 1
                
                # Extract title from content if not in properties
                if metadata['title'] == file_path.stem and text_content:
                    first_line = text_content[0]
                    if len(first_line) > 5:  # Skip very short lines
                        metadata['title'] = first_line[:100]  # Limit title length
                
            elif file_ext in ['.doc']:
                # For .doc files, we can't easily extract without additional libraries
                # Provide basic file info
                metadata['has_text_content'] = True  # Assume Office docs have text
                metadata['word_count'] = 0
                metadata['page_count'] = 1
                metadata['author'] = ''
                metadata['title'] = file_path.stem
                metadata['subject'] = ''
                metadata['keywords'] = ''
                metadata['category'] = ''
                metadata['comments'] = ''
                metadata['created'] = ''
                metadata['modified'] = ''
                metadata['last_modified_by'] = ''
                
            else:
                # Other office formats (xlsx, pptx, etc.)
                metadata['has_text_content'] = True  # Assume Office docs have text
                metadata['word_count'] = 0
                metadata['page_count'] = 1
                metadata['author'] = ''
                metadata['title'] = file_path.stem
                metadata['subject'] = ''
                metadata['keywords'] = ''
                metadata['category'] = ''
                metadata['comments'] = ''
                metadata['created'] = ''
                metadata['modified'] = ''
                metadata['last_modified_by'] = ''

        except Exception as e:
            # If document processing fails, provide basic info
            metadata['has_text_content'] = False
            metadata['word_count'] = 0
            metadata['page_count'] = 1
            metadata['author'] = ''
            metadata['title'] = file_path.stem
            metadata['error'] = str(e)

        return metadata

    def _parse_frontmatter(self, content: str) -> tuple[Dict[str, Any], str]:
        """Parse YAML frontmatter from markdown content."""
        frontmatter = {}
        body = content

        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                try:
                    frontmatter = yaml.safe_load(parts[1]) or {}
                    body = parts[2].strip()
                except yaml.YAMLError:
                    pass  # Invalid YAML, keep empty frontmatter

        return frontmatter, body

    def _extract_tags(self, frontmatter: Dict[str, Any], body: str) -> List[str]:
        """Extract tags from frontmatter and body."""
        tags = set()

        # From frontmatter
        fm_tags = frontmatter.get('tags', [])
        if isinstance(fm_tags, str):
            tags.update(t.strip() for t in fm_tags.split(',') if t.strip())
        elif isinstance(fm_tags, list):
            tags.update(str(t).strip() for t in fm_tags if str(t).strip())

        # From body (#tag format)
        body_tags = re.findall(r'(?<!\w)#([a-zA-Z0-9_/-]+)', body)
        tags.update(body_tags)

        return sorted(list(tags))

    def _extract_wiki_links(self, body: str) -> List[str]:
        """Extract Obsidian wiki links from content."""
        # Match [[link]] or [[link|alias]]
        links = re.findall(r'\[\[([^\]]+)\]\]', body)
        # Remove aliases, keep just the link targets
        return [link.split('|')[0].strip() for link in links]

    def get_description_prompt_template(self) -> str:
        return """You are a technical documentation assistant. Generate a one-sentence description and category for a document based on its content and metadata.

Available categories (choose ONE):
- research_papers: Academic papers, research articles, and scholarly documents
- business_docs: Business plans, reports, proposals, and corporate documents
- legal_documents: Contracts, agreements, legal briefs, and compliance documents
- educational_materials: Textbooks, course materials, tutorials, and learning resources
- technical_docs: API documentation, manuals, specifications, and technical guides
- personal_docs: Personal letters, journals, memoirs, and private documents
- reports_presentations: Reports, presentations, whitepapers, and analytical documents
- utilities_misc: Forms, templates, checklists, and miscellaneous documents

Document Metadata:
File Type: {file_extension}
Word Count: {word_count}
Page Count: {page_count}
Author: {author}
Title: {title}

Content Sample:
---
{content}
---

Generate a JSON response with:
1. "description": A single-sentence description (max 150 characters) that captures the document's core purpose and content. Be concise and technical.
2. "category": ONE category from the list above that best matches this document.

Example format:
{"description": "Comprehensive business plan outlining market strategy and financial projections for startup launch", "category": "business_docs"}

JSON Response:"""

    def get_example_descriptions(self) -> List[str]:
        return [
            "Comprehensive research paper on machine learning algorithms and their applications in data science",
            "Business proposal outlining market strategy and financial projections for new product launch",
            "Legal contract specifying terms and conditions for software licensing agreement",
            "Educational tutorial explaining advanced mathematics concepts for computer science students",
            "Technical specification document detailing API endpoints and data structures",
            "Personal journal documenting learning experiences and professional development goals",
            "Annual business report summarizing company performance and strategic initiatives"
        ]

    def get_content_for_description(self, item: CollectionItem) -> str:
        """
        Extract content from document for LLM description generation.
        Returns first 3000 chars of readable content.
        """
        file_path = Path(item.path)
        file_ext = file_path.suffix.lower()

        try:
            if file_ext in ['.txt', '.md', '.tex']:
                # Text-based documents
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return content[:3000]
                
            elif file_ext == '.pdf' and PYPDF2_AVAILABLE:
                # Extract text from PDF
                try:
                    with open(file_path, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        text_content = ''
                        
                        # Extract text from first few pages
                        max_pages = min(3, len(pdf_reader.pages))
                        for page_num in range(max_pages):
                            try:
                                page = pdf_reader.pages[page_num]
                                text_content += page.extract_text() + '\n'
                                if len(text_content) > 3000:
                                    break
                            except Exception:
                                continue
                        
                        return text_content[:3000] if text_content.strip() else f"PDF document: {file_path.stem}"
                except Exception:
                    return f"PDF document: {file_path.stem}"
                    
            elif file_ext == '.docx' and PYTHON_DOCX_AVAILABLE:
                # Extract text from Word document
                try:
                    doc = Document(file_path)
                    text_content = []
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
                            # Stop if we have enough content
                            if len('\n'.join(text_content)) > 3000:
                                break
                    
                    full_text = '\n'.join(text_content)
                    return full_text[:3000] if full_text.strip() else f"Word document: {file_path.stem}"
                except Exception:
                    return f"Word document: {file_path.stem}"
                    
            elif file_ext in ['.doc']:
                # For .doc files, we can't easily extract text
                return f"Word document: {file_path.stem}"
                
            else:
                # Other document types
                return f"Document: {file_path.stem}"
                
        except Exception:
            return f"Document: {file_path.stem}"


# Register plugin on import
PluginRegistry.register(
    name="documents",
    scanner_class=DocumentsScanner,
    version="1.0.0",
    description="Scanner for document collections with metadata extraction from PDFs, Office docs, and text files"
)